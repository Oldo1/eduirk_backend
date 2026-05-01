"""
doc_extractor.py — извлечение текста из документов

Поддерживаемые форматы:
    .pdf   — нативный текст + OCR-fallback для страниц-сканов
    .docx  — параграфы и таблицы через python-docx
    .doc   — конвертация в .docx через Word COM (pywin32) или mammoth,
             затем обработка как .docx

Зависимости:
    pip install pypdf python-docx
    pip install pymupdf surya-ocr pillow     # для OCR (опционально)
    pip install pywin32                       # для .doc через Word (опционально)
    pip install mammoth                       # для .doc без Word (опционально)
"""

from __future__ import annotations

import io
import logging
import re
import tempfile
from pathlib import Path
from typing import Optional

from config import PDF_MIN_PAGE_CHARS
from ocr_engine import is_ocr_available, ocr_pdf_pages
from ocr_cache import get_cached, save_cached

logger = logging.getLogger("doc_extractor")


# ─────────────────────────────────────────────────────────────────────────────
#  PDF
# ─────────────────────────────────────────────────────────────────────────────

def extract_pdf(pdf_bytes: bytes) -> str:
    """
    Извлекает текст из PDF постранично.
    Сначала пытается нативное извлечение, затем ОДНИМ БАТЧЕМ прогоняет
    через OCR все страницы-сканы. Батч-OCR в 2-3× быстрее поштучного
    благодаря параллельной детекции внутри Surya.

    Returns:
        Полный текст документа или пустую строку при ошибке.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("[pdf] pypdf не установлен: pip install pypdf")
        return ""

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        total  = len(reader.pages)

        # 1. Нативное извлечение + список страниц-кандидатов на OCR
        native_texts: dict[int, str] = {}
        scan_pages:   list[int]      = []

        for i, page in enumerate(reader.pages):
            raw = re.sub(r'\s+', ' ', page.extract_text() or "").strip()
            if len(raw) >= PDF_MIN_PAGE_CHARS:
                native_texts[i] = raw
            else:
                scan_pages.append(i)

        # 2. Батч-OCR для всех сканов одним вызовом
        ocr_texts: dict[int, str] = {}
        if scan_pages:
            if is_ocr_available():
                logger.info(f"[pdf] Страниц-сканов: {len(scan_pages)}/{total}, запускаю батч-OCR")
                ocr_texts = ocr_pdf_pages(pdf_bytes, scan_pages)
                # Оставляем только те, где удалось что-то извлечь
                ocr_texts = {
                    idx: txt for idx, txt in ocr_texts.items()
                    if len(txt) >= PDF_MIN_PAGE_CHARS
                }
                logger.info(f"[pdf] OCR извлёк текст на {len(ocr_texts)}/{len(scan_pages)} страницах")
            else:
                logger.debug(f"[pdf] {len(scan_pages)} страниц-сканов, но OCR недоступен")

        # 3. Собираем результат в исходном порядке страниц
        parts: list[str] = []
        for i in range(total):
            if i in native_texts:
                parts.append(native_texts[i])
            elif i in ocr_texts:
                parts.append(ocr_texts[i])

        return "\n\n".join(parts)

    except Exception as e:
        logger.error(f"[pdf] Ошибка при чтении: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX
# ─────────────────────────────────────────────────────────────────────────────

def extract_docx(docx_bytes: bytes) -> str:
    """
    Извлекает текст из .docx:
      — параграфы (основной текст, заголовки)
      — таблицы (строка → «ячейка | ячейка | ячейка»)

    Returns:
        Полный текст документа или пустую строку при ошибке.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("[docx] python-docx не установлен: pip install python-docx")
        return ""

    try:
        doc   = DocxDocument(io.BytesIO(docx_bytes))
        lines: list[str] = []

        # Параграфы
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Таблицы
        for table_idx, table in enumerate(doc.tables):
            seen_rows: set[str] = set()   # защита от дублей объединённых ячеек
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if not cells:
                    continue
                row_str = " | ".join(cells)
                if row_str not in seen_rows:
                    lines.append(row_str)
                    seen_rows.add(row_str)

        text = "\n".join(lines)
        # Убираем тройные и более переносы строк
        return re.sub(r'\n{3,}', '\n\n', text).strip()

    except Exception as e:
        logger.error(f"[docx] Ошибка при чтении: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
#  DOC → DOCX конвертация
# ─────────────────────────────────────────────────────────────────────────────

def _convert_via_word_com(doc_bytes: bytes, filename: str) -> Optional[bytes]:
    """
    Конвертирует .doc → .docx через Microsoft Word (COM-автоматизация).
    Требует: pip install pywin32  и  установленный Microsoft Word.
    """
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return None   # pywin32 не установлен — тихо пропускаем

    try:
        pythoncom.CoInitialize()
        word = None

        with tempfile.TemporaryDirectory() as tmp:
            doc_path  = Path(tmp) / filename
            docx_path = doc_path.with_suffix(".docx")
            doc_path.write_bytes(doc_bytes)

            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                document = word.Documents.Open(str(doc_path.resolve()))
                try:
                    # 16 = wdFormatXMLDocument (.docx)
                    document.SaveAs2(str(docx_path.resolve()), FileFormat=16)
                finally:
                    document.Close(SaveChanges=False)
            finally:
                if word is not None:
                    word.Quit()
                pythoncom.CoUninitialize()

            if docx_path.exists():
                return docx_path.read_bytes()

    except Exception as e:
        logger.debug(f"[doc] Word COM ошибка: {e}")

    return None


def _convert_via_mammoth(doc_bytes: bytes) -> Optional[bytes]:
    """
    Извлекает текст из .doc через mammoth и упаковывает в .docx.
    Требует: pip install mammoth python-docx
    Работает без Microsoft Word и LibreOffice.
    """
    try:
        import mammoth
    except ImportError:
        return None   # mammoth не установлен — тихо пропускаем

    try:
        from docx import Document as DocxDocument

        result   = mammoth.extract_raw_text(io.BytesIO(doc_bytes))
        raw_text = result.value.strip()

        if not raw_text:
            logger.warning("[doc] mammoth вернул пустой текст")
            return None

        # Упаковываем в .docx чтобы переиспользовать extract_docx
        new_doc = DocxDocument()
        for line in raw_text.splitlines():
            line = line.strip()
            if line:
                new_doc.add_paragraph(line)

        buf = io.BytesIO()
        new_doc.save(buf)
        return buf.getvalue()

    except Exception as e:
        logger.debug(f"[doc] mammoth ошибка: {e}")
        return None


def convert_doc_to_docx(doc_bytes: bytes, filename: str) -> Optional[bytes]:
    """
    Конвертирует .doc → .docx, пробуя стратегии по порядку:
      1. Microsoft Word COM (pywin32) — лучшее качество, нужен Word
      2. mammoth                       — без Word, только текст

    Returns:
        Байты .docx или None если конвертация не удалась.
    """
    logger.info(f"[doc] Конвертирую {filename} → .docx...")

    # Стратегия 1: Word COM
    result = _convert_via_word_com(doc_bytes, filename)
    if result:
        logger.info(f"[doc] Конвертировано через Microsoft Word")
        return result

    # Стратегия 2: mammoth
    result = _convert_via_mammoth(doc_bytes)
    if result:
        logger.info(f"[doc] Конвертировано через mammoth")
        return result

    logger.warning(
        f"[doc] Не удалось конвертировать {filename}.\n"
        "      Установите одно из: pip install pywin32  |  pip install mammoth"
    )
    return None


# ─────────────────────────────────────────────────────────────────────────────
#  Диспетчер: выбор метода по расширению
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Извлекает текст из файла, выбирая метод по расширению.
    Использует файловый кэш по SHA-256 содержимого — при повторном
    извлечении того же файла возвращает результат мгновенно, без OCR.

    Args:
        file_bytes: содержимое файла в байтах
        filename:   имя файла (нужно для определения расширения и конвертации .doc)

    Returns:
        Извлечённый текст или пустую строку.
    """
    # Проверяем кэш до любой тяжёлой работы
    cached = get_cached(file_bytes)
    if cached is not None:
        return cached

    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = extract_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_docx(file_bytes)
    elif ext == ".doc":
        docx_bytes = convert_doc_to_docx(file_bytes, filename)
        text = extract_docx(docx_bytes) if docx_bytes else ""
    else:
        logger.warning(f"[extractor] Неподдерживаемый формат: {ext} ({filename})")
        text = ""

    # Сохраняем только непустой результат
    save_cached(file_bytes, text)
    return text
